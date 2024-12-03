from docx import Document

class ReportGenerator:
    @staticmethod
    def create_word_report(data, file_path):
        doc = Document()
        doc.add_heading("Analysis Report", level=1)

        print(data)
        
        for record in data:
            doc.add_paragraph(f"Transformador: {data[-1]['transformer_identification.transformer_id']}")
            doc.add_paragraph(f"Método: {record['method']}")
            doc.add_paragraph(f"Resultado: {record['result']}")
            doc.add_paragraph(f"Data da Análise: {record['analysis_date']}")
            doc.add_paragraph("")

        doc.save(file_path)
